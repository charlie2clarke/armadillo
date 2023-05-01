import docx
from models.document import Document


class DocumentDao:
    ''' Data access object for interacting with word documents '''

    def __init__(self, file_path):
        ''' Initialises the document instance with the file path '''
        self.document = Document(file_path)

        self.initialise_document()

    def initialise_document(self):
        ''' Initialises the document instance with the tables, paragraphs and words '''
        self.document.tables = self.document.doc.tables
        self.document.paras = list(self.set_paragraphs())
        self.document.words = list(self.set_words())

    def set_paragraphs(self):
        ''' Yields all the paragraphs in the document, skipping headings, 
            empty paragraphs, figure captions, and the appendices '''
        counter = 0

        while counter < len(self.document.doc.paragraphs): 
            paragraph = self.document.doc.paragraphs[counter]
            counter += 1

            if paragraph.text.startswith('Appendix') or paragraph.text.startswith('Appendices'):
                break

            if not paragraph.style.name.startswith('Heading') and not paragraph.style.name.startswith('Title') and not paragraph.style.name.startswith('Subtitle') and \
                paragraph.text != '' and not paragraph.text.startswith('Figure'):
                yield paragraph


        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text != '':
                            yield paragraph

    paragraph_count = 0

    def set_words(self):
        ''' Yields words from the paragraphs, excluding full stops and dashes '''
        for paragraph in self.document.paras:
            self.paragraph_count += 1
            print('Paragraph ' + str(self.paragraph_count) + ' style:' + str(paragraph.style.name))
            paragraph = paragraph.text
            words = paragraph.split()
            for word in words:
                print(word) 
                if word != ('.' or '–'):
                    yield word

    def print_word_count(self):
        ''' Prints total word count of document '''
        print(len(self.document.words))
